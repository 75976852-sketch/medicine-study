#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取Word文档"中药分类.docx"
"""

try:
    from docx import Document
    
    doc_path = '中药分类.docx'
    print(f"正在读取: {doc_path}")
    print("=" * 70)
    
    doc = Document(doc_path)
    
    print(f"文档共有 {len(doc.paragraphs)} 个段落\n")
    
    # 读取所有段落
    all_text = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append(text)
            if i < 50:  # 只打印前50个段落
                print(f"{i+1}. {text[:150]}")
    
    print(f"\n... (共 {len(all_text)} 个非空段落)")
    print("=" * 70)
    
    # 将内容保存到文本文件
    with open('中药分类_提取.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(all_text))
    
    print(f"\n内容已保存到: 中药分类_提取.txt")
    
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip3 install python-docx")
except FileNotFoundError:
    print(f"错误: 找不到文件 '中药分类.docx'")
except Exception as e:
    print(f"读取时出错: {e}")

